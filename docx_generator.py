"""
Word 문서 생성 모듈
python-docx를 사용하여 국어 시험지 레이아웃을 재현
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List


def create_document() -> Document:
    """
    새 Word 문서를 생성하고 기본 설정을 적용합니다.
    
    Returns:
        Document: 생성된 문서 객체
    """
    doc = Document()
    
    # 여백 설정 (A4 기준)
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 세로
    section.page_width = Inches(8.27)    # A4 가로
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    
    return doc


def create_two_column_layout(doc: Document):
    """
    2단 레이아웃을 설정합니다.
    
    Args:
        doc: 문서 객체
    """
    section = doc.sections[0]
    
    # 2단 설정
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if not cols:
        cols_elem = OxmlElement('w:cols')
        sectPr.append(cols_elem)
    else:
        cols_elem = cols[0]
    
    cols_elem.set(qn('w:num'), '2')
    cols_elem.set(qn('w:space'), '720')  # 단 간격 (1/20 inch 단위)


def apply_korean_font(doc: Document, font_name: str = "바탕"):
    """
    전체 문서에 한글 폰트를 적용합니다.
    
    Args:
        doc: 문서 객체
        font_name: 폰트 이름 (기본값: 바탕)
    """
    # 기본 스타일에 폰트 적용
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(10)
    
    # 동아시아 폰트 설정
    rPr = style.element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.insert(0, rPr)
    
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    
    rFonts.set(qn('w:eastAsia'), font_name)


def add_passage(doc: Document, passage_text: str):
    """
    지문을 추가합니다.
    
    Args:
        doc: 문서 객체
        passage_text: 지문 텍스트
    """
    if passage_text:
        p = doc.add_paragraph(passage_text)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2


def add_box_content(doc: Document, box_text: str):
    """
    보기 박스를 테이블로 표현합니다.
    
    Args:
        doc: 문서 객체
        box_text: 보기 내용
    """
    if not box_text:
        return
    
    # 1행 1열 테이블 생성
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    # 테두리 설정
    cell = table.rows[0].cells[0]
    cell.text = box_text
    
    # 셀 내부 여백
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), '100')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    # 표 뒤 간격
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def add_question(doc: Document, question_text: str):
    """
    문제를 추가합니다.
    
    Args:
        doc: 문서 객체
        question_text: 문제 텍스트
    """
    if question_text:
        p = doc.add_paragraph(question_text)
        run = p.runs[0] if p.runs else p.add_run()
        run.bold = True
        p.paragraph_format.space_after = Pt(6)


def add_options_flexible(doc: Document, options: List[str]):
    """
    선지를 유연하게 배치합니다.
    
    Args:
        doc: 문서 객체
        options: 선택지 리스트
    """
    if not options:
        return
    
    # 평균 길이 계산
    avg_length = sum(len(opt) for opt in options) / len(options) if options else 0
    
    if avg_length < 30:  # 짧은 선지 → 테이블로 가로 나열
        # 2열 테이블 (한 행에 2개씩)
        num_rows = (len(options) + 1) // 2
        table = doc.add_table(rows=num_rows, cols=2)
        
        for i, option in enumerate(options):
            row_idx = i // 2
            col_idx = i % 2
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = option
        
        # 테이블 스타일 제거 (테두리 없음)
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
    
    else:  # 긴 선지 → 세로 나열
        for option in options:
            p = doc.add_paragraph(option)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(3)
    
    # 문항 끝 간격
    doc.add_paragraph()


def generate_word_document(data: Dict, output_path: str, use_two_columns: bool = True) -> bool:
    """
    JSON 데이터로부터 Word 문서를 생성합니다.
    
    Args:
        data: JSON 데이터 (parse_to_json 결과)
        output_path: 출력 파일 경로
        use_two_columns: 2단 레이아웃 사용 여부
        
    Returns:
        bool: 생성 성공 여부
    """
    try:
        doc = create_document()
        apply_korean_font(doc)
        
        if use_two_columns:
            create_two_column_layout(doc)
        
        items = data.get("items", [])
        
        for item in items:
            # 지문
            add_passage(doc, item.get("passage", ""))
            
            # 보기 박스
            add_box_content(doc, item.get("box_content", ""))
            
            # 문제
            add_question(doc, item.get("question", ""))
            
            # 선지
            add_options_flexible(doc, item.get("options", []))
        
        # 문서 저장
        doc.save(output_path)
        print(f"문서 생성 완료: {output_path}")
        return True
        
    except Exception as e:
        print(f"문서 생성 오류: {e}")
        return False
